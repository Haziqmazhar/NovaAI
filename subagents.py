"""
subagents.py — specialized sub-agents the orchestrator (brain.py) can
delegate to. Each sub-agent gets its own narrow, tool-less GPT call and
zero direct system access — the only thing it can touch is what its own
function explicitly reads, inside an explicit whitelist check. Extend
this file with one function per new sub-agent capability.
"""

import csv
import os
import urllib.parse
import urllib.request

MAX_CHARS = 20_000
MAX_FILES_PER_CALL = 5
MAX_LISTED_FILES = 50
REQUEST_TIMEOUT_SECONDS = 30.0
DOWNLOAD_TIMEOUT_SECONDS = 30.0
MAX_DOWNLOAD_BYTES = 50 * 1024 * 1024  # 50MB safety cap

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".xlsx"}

# Refuse to save these — download_file is the project's only *write* path,
# so this is defense-in-depth against Nova being tricked into planting
# something executable on disk, even though open_item can't launch it
# anyway (it only launches apps/folders already in config.json).
BLOCKED_DOWNLOAD_EXTENSIONS = {
    ".exe", ".bat", ".cmd", ".com", ".msi", ".ps1", ".vbs", ".js",
    ".jar", ".scr", ".pif", ".apk", ".sh", ".dll", ".lnk", ".msc",
}
_INVALID_FILENAME_CHARS = '<>:"/\\|?*'

DOCUMENT_AGENT_SYSTEM_PROMPT = (
    "You are a document analysis sub-agent. You are given the text "
    "contents of one or more files and an instruction. Do only what the "
    "instruction asks, using only the given content — you have no tools "
    "and no other context. If asked to compare or cross-reference files, "
    "use the '=== filename ===' section headers to tell them apart. Be "
    "concise."
)


def _resolve_whitelisted_file(config: dict, file_name: str):
    """Find file_name inside one of config['folders'], refusing to leave
    the whitelisted folder even via '..' or symlink tricks. Returns the
    resolved absolute path, or None if not found / not allowed."""
    folders = config.get("folders", {})
    target_basename = os.path.basename(file_name)

    for folder_path in folders.values():
        root = os.path.realpath(os.path.expandvars(os.path.expanduser(folder_path)))
        if not os.path.isdir(root):
            continue
        candidate = os.path.realpath(os.path.join(root, target_basename))
        if os.path.commonpath([root, candidate]) == root and os.path.isfile(candidate):
            return candidate
    return None


def _resolve_whitelisted_folders(config: dict, folder_name: str = None):
    """Match folder_name against config['folders'] (case-insensitive
    exact/substring, same style as actions.py::find_target). Returns a
    dict of {key: root_path} — either the single matched folder, or every
    whitelisted folder if folder_name is None. Returns {} if folder_name
    was given but didn't match anything."""
    folders = config.get("folders", {})
    resolved = {}
    for key, folder_path in folders.items():
        root = os.path.realpath(os.path.expandvars(os.path.expanduser(folder_path)))
        if os.path.isdir(root):
            resolved[key] = root

    if folder_name is None:
        return resolved

    needle = folder_name.strip().lower()
    for key, root in resolved.items():
        if needle in key.lower() or key.lower() in needle:
            return {key: root}
    return {}


def list_documents(config: dict, folder_name: str = None) -> str:
    """List readable files (SUPPORTED_EXTENSIONS) inside a whitelisted
    folder, or across all whitelisted folders if folder_name is None.
    Never raises — returns a plain-language message on any failure."""
    folders = _resolve_whitelisted_folders(config, folder_name)
    if not folders:
        known = ", ".join(sorted(config.get("folders", {}).keys())) or "none configured"
        return f"'{folder_name}' isn't one of the whitelisted folders. Folders available: {known}."

    sections = []
    for key, root in folders.items():
        try:
            entries = os.listdir(root)
        except OSError as e:
            sections.append(f"{key}: could not list ({e})")
            continue

        readable = sorted(e for e in entries if os.path.splitext(e)[1].lower() in SUPPORTED_EXTENSIONS)
        skipped = len(entries) - len(readable)

        if not readable:
            line = f"{key}: no readable documents"
        else:
            shown = readable[:MAX_LISTED_FILES]
            more = len(readable) - len(shown)
            line = f"{key}: " + ", ".join(shown)
            if more:
                line += f" (+{more} more)"
        if skipped:
            line += f" [{skipped} other file(s) with unsupported formats omitted]"
        sections.append(line)

    return "\n".join(sections)


def _safe_filename(name: str) -> str:
    """Strip any path components and characters invalid on Windows, plus
    leading/trailing dots/spaces (so e.g. '..' can never survive as a
    name). Returns '' if nothing usable is left."""
    name = os.path.basename(name or "").strip()
    name = "".join(c for c in name if c not in _INVALID_FILENAME_CHARS and ord(c) >= 32)
    return name.strip(" .")


def _unique_path(root: str, file_name: str) -> str:
    """Never overwrite an existing file — append ' (1)', ' (2)', etc.
    until a free name is found."""
    base, ext = os.path.splitext(file_name)
    candidate = os.path.join(root, file_name)
    n = 1
    while os.path.exists(candidate):
        candidate = os.path.join(root, f"{base} ({n}){ext}")
        n += 1
    return candidate


def download_file(config: dict, url: str, folder_name: str, file_name: str = None) -> str:
    """Download `url` into one whitelisted folder. Only http(s) URLs are
    allowed, executable/script extensions are refused, size is capped at
    MAX_DOWNLOAD_BYTES, and an existing file is never overwritten. Never
    raises — returns a plain-language result either way."""
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https"):
        return f"Refusing to download '{url}' — only http/https URLs are allowed."

    folders = _resolve_whitelisted_folders(config, folder_name)
    if not folders:
        known = ", ".join(sorted(config.get("folders", {}).keys())) or "none configured"
        return f"'{folder_name}' isn't one of the whitelisted folders. Folders available: {known}."
    key, root = next(iter(folders.items()))

    candidate_name = _safe_filename(file_name) or _safe_filename(os.path.basename(parsed.path))
    if not candidate_name:
        candidate_name = "download"

    ext = os.path.splitext(candidate_name)[1].lower()
    if ext in BLOCKED_DOWNLOAD_EXTENSIONS:
        return f"Refusing to download '{candidate_name}' — executable/script files aren't allowed."

    dest_path = _unique_path(root, candidate_name)
    part_path = dest_path + ".part"

    written = 0
    try:
        request = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 (Nova Assistant)"})
        with urllib.request.urlopen(request, timeout=DOWNLOAD_TIMEOUT_SECONDS) as response:
            content_length = response.headers.get("Content-Length")
            if content_length and int(content_length) > MAX_DOWNLOAD_BYTES:
                mb = int(content_length) // (1024 * 1024)
                return f"Refusing to download — file is {mb}MB, over the {MAX_DOWNLOAD_BYTES // (1024 * 1024)}MB limit."

            with open(part_path, "wb") as f:
                while True:
                    chunk = response.read(65536)
                    if not chunk:
                        break
                    written += len(chunk)
                    if written > MAX_DOWNLOAD_BYTES:
                        raise ValueError(f"exceeded the {MAX_DOWNLOAD_BYTES // (1024 * 1024)}MB download limit")
                    f.write(chunk)
    except Exception as e:
        if os.path.exists(part_path):
            os.remove(part_path)
        return f"Failed to download '{url}': {e}"

    os.replace(part_path, dest_path)
    size_kb = written / 1024
    return f"Downloaded '{os.path.basename(dest_path)}' ({size_kb:.0f} KB) to the '{key}' folder."


def _extract_text(path: str, max_chars: int = MAX_CHARS) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        from docx import Document

        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
    elif ext == ".csv":
        with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
            rows = csv.reader(f)
            text = "\n".join("\t".join(row) for row in rows)
    elif ext == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"== Sheet: {sheet.title} ==")
            for row in sheet.iter_rows(values_only=True):
                parts.append("\t".join("" if v is None else str(v) for v in row))
        text = "\n".join(parts)
    else:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

    if len(text) > max_chars:
        return text[:max_chars] + "\n[...truncated...]"
    return text


def run_document_agent(config: dict, client, model: str, instruction: str, file_names: list) -> str:
    """Read one or more whitelisted files and have a single scoped GPT
    call act on them per `instruction` (e.g. summarize, compare). Returns
    the result text, or a plain-language error string on any failure —
    never raises, so a bad file/instruction can't crash the orchestrator's
    turn."""
    notes = []
    if len(file_names) > MAX_FILES_PER_CALL:
        notes.append(
            f"Only the first {MAX_FILES_PER_CALL} files were used; "
            f"skipped: {', '.join(file_names[MAX_FILES_PER_CALL:])}."
        )
        file_names = file_names[:MAX_FILES_PER_CALL]

    resolved = []  # list of (file_name, path)
    missing = []
    for file_name in file_names:
        path = _resolve_whitelisted_file(config, file_name)
        if path:
            resolved.append((file_name, path))
        else:
            missing.append(file_name)

    if missing:
        notes.append(
            f"Not found in any whitelisted folder, so excluded: {', '.join(missing)}."
        )

    if not resolved:
        return " ".join(notes) or "No valid files were given."

    per_file_budget = max(3000, MAX_CHARS // len(resolved))
    sections = []
    for file_name, path in resolved:
        try:
            text = _extract_text(path, max_chars=per_file_budget)
        except Exception as e:
            notes.append(f"Could not read '{file_name}': {e}")
            continue
        if not text.strip():
            notes.append(f"'{file_name}' appears to be empty, or its text couldn't be extracted.")
            continue
        sections.append(f"=== {file_name} ===\n{text}")

    if not sections:
        return " ".join(notes) or "None of the given files could be read."

    prompt_parts = []
    if notes:
        prompt_parts.append("Note: " + " ".join(notes))
    prompt_parts.append(f"Instruction: {instruction}")
    prompt_parts.append("File contents:\n\n" + "\n\n".join(sections))
    prompt = "\n\n".join(prompt_parts)

    last_error = None
    for _attempt in range(2):  # one retry on a transient API failure
        try:
            response = client.with_options(timeout=REQUEST_TIMEOUT_SECONDS).responses.create(
                model=model,
                instructions=DOCUMENT_AGENT_SYSTEM_PROMPT,
                input=[{"role": "user", "content": prompt}],
            )
            return (response.output_text or "").strip()
        except Exception as e:
            last_error = e

    names = ", ".join(name for name, _ in resolved)
    return f"The document agent failed to process '{names}': {last_error}"
